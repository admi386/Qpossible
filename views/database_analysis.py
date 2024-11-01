import asyncio
import streamlit as st
import pandas as pd
from sqlalchemy import create_engine
from io import BytesIO
from docx import Document
from settings_manager import load_settings, save_settings
from db_settings_manager import load_db_settings, save_db_settings

# Load settings from a file
if "llm_settings" not in st.session_state:
    st.session_state.llm_settings = load_settings()

if "db_settings" not in st.session_state:
    st.session_state.db_settings = load_db_settings()

# Initialize state for the current page using a unique key
page_key = "db_analysis"
if f"{page_key}_report" not in st.session_state:
    st.session_state[f"{page_key}_report"] = None  # No report initially

if f"{page_key}_llm_response" not in st.session_state:
    st.session_state[f"{page_key}_llm_response"] = None  # No LLM response initially

if f"{page_key}_metadata_description" not in st.session_state:
    st.session_state[f"{page_key}_metadata_description"] = None  # No database metadata description initially

# Universal wrapper for working with LLM
class LLMWrapper:
    def __init__(self, settings):
        self.settings = settings
        self.llm = self._initialize_llm()

    def _initialize_llm(self):
        llm_type = self.settings["llm_type"]
        if llm_type == "ollama":
            from langchain_community.llms import Ollama
            return Ollama(
                model=self.settings["model"],
                temperature=self.settings["temperature"]
            )
        elif llm_type == "groq":
            from groq import Groq
            return Groq(
                api_key=self.settings["api_key"]
            )
        elif llm_type == "openai":
            from openai import OpenAI
            return OpenAI(
                api_key=self.settings["api_key"],
                temperature=self.settings["temperature"],
                max_tokens=self.settings["max_tokens"]
            )
        else:
            raise ValueError(f"Unsupported LLM type: {llm_type}")

    async def generate_response_async(self, prompt):
        if self.settings["llm_type"] == "groq":
            chat_completion = await self.llm.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=self.settings["model"],
            )
            return chat_completion.choices[0].message.content
        else:
            return await asyncio.to_thread(self.llm, prompt)

# Initialize LLM
llm_wrapper = LLMWrapper(st.session_state.llm_settings)

# Streamlit Interface
st.title("Database Analysis")
st.write("This tool analyzes the structure of a database and suggests types of analysis that can be performed.")

# Function to connect to the database
def connect_to_database():
    db_settings = st.session_state.db_settings
    db_type = db_settings["db_type"]
    try:
        if db_type == "PostgreSQL":
            engine = create_engine(f"postgresql://{db_settings['username']}:{db_settings['password']}@{db_settings['host']}:{db_settings['port']}/{db_settings['db_name']}")
        elif db_type == "MySQL":
            engine = create_engine(f"mysql+pymysql://{db_settings['username']}:{db_settings['password']}@{db_settings['host']}:{db_settings['port']}/{db_settings['db_name']}")
        elif db_type == "SQLite":
            engine = create_engine(f"sqlite:///{db_settings['db_name']}")
        elif db_type == "SQL Server":
            engine = create_engine(f"mssql+pyodbc://{db_settings['username']}:{db_settings['password']}@{db_settings['host']}/{db_settings['db_name']}?driver=SQL+Server")
        else:
            st.error("Unsupported database type")
            return None
        return engine.connect()
    except Exception as e:
        st.error(f"Database connection error: {e}")
        return None

# Function to get database metadata, including relationships between tables
def get_database_metadata(connection):
    db_settings = st.session_state.db_settings
    db_type = db_settings["db_type"]
    metadata = {}
    try:
        if db_type == "SQLite":
            tables_query = "SELECT name FROM sqlite_master WHERE type='table';"
            tables_result = pd.read_sql_query(tables_query, connection)
            tables_result = tables_result.rename(columns={'name': 'table_name'})
        else:
            tables_query = "SELECT table_name FROM information_schema.tables WHERE table_schema = 'public';"
            tables_result = pd.read_sql_query(tables_query, connection)
        
        if 'table_name' not in tables_result.columns:
            st.error("Column 'table_name' not found in tables result.")
            return metadata

        for table in tables_result['table_name']:
            # Retrieve column information
            if db_type == "SQLite":
                columns_query = f"PRAGMA table_info({table});"
                columns_result = pd.read_sql_query(columns_query, connection)
                columns = [{'column_name': row['name'], 'data_type': row['type']} for _, row in columns_result.iterrows()]
            else:
                columns_query = f"SELECT column_name, data_type FROM information_schema.columns WHERE table_name = '{table}';"
                columns_result = pd.read_sql_query(columns_query, connection)
                columns = columns_result.to_dict(orient='records')

            # Retrieve foreign key information
            if db_type == "SQLite":
                foreign_keys_query = f"PRAGMA foreign_key_list({table});"
                foreign_keys_result = pd.read_sql_query(foreign_keys_query, connection)
                foreign_keys = [{'source_column': row['from'], 'target_table': row['table'], 'target_column': row['to']} for _, row in foreign_keys_result.iterrows()]
            else:
                foreign_keys_query = f"""
                SELECT 
                    kcu.column_name AS source_column,
                    ccu.table_name AS target_table,
                    ccu.column_name AS target_column 
                FROM 
                    information_schema.table_constraints AS tc 
                    JOIN information_schema.key_column_usage AS kcu
                      ON tc.constraint_name = kcu.constraint_name
                      AND tc.table_schema = kcu.table_schema
                    JOIN information_schema.constraint_column_usage AS ccu
                      ON ccu.constraint_name = tc.constraint_name
                      AND ccu.table_schema = tc.table_schema
                WHERE tc.constraint_type = 'FOREIGN KEY' AND tc.table_name='{table}';
                """
                foreign_keys_result = pd.read_sql_query(foreign_keys_query, connection)
                foreign_keys = foreign_keys_result.to_dict(orient='records')

            metadata[table] = {'columns': columns, 'foreign_keys': foreign_keys}
    except Exception as e:
        st.error(f"Error retrieving database metadata: {e}")
    return metadata

# Function to create a report in DOCX format
def save_report_to_docx(report_text):
    doc = Document()
    doc.add_heading('Database Overview Report', 0)
    doc.add_paragraph(report_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to analyze the database and provide recommendations
async def analyze_database(connection):
    # Retrieve database metadata
    metadata = get_database_metadata(connection)
    
    # Formulate a prompt for LLM describing the data structure and table relationships
    metadata_description = "\n".join(
        [f"Table: {table}\nColumns: {', '.join([col['column_name'] + ' (' + col['data_type'] + ')' for col in info['columns']])}"
         + "\nForeign Keys: " + (', '.join(['{} -> {}.{}'.format(fk['source_column'], fk['target_table'], fk['target_column']) for fk in info['foreign_keys']]) if info['foreign_keys'] else 'None')
         for table, info in metadata.items()]
    )
    
    # Save metadata description to session state
    st.session_state[f"{page_key}_metadata_description"] = metadata_description
    
    prompt = f"""
    You are an AI assistant analyzing a database. The database structure is as follows:
    
    {metadata_description}
    
    1. Summarize the data contained in this database.
    2. Suggest possible types of analysis that can be performed based on this data.
    3. Describe the potential benefits of conducting these analyses.
    
    Consider relationships between tables (Foreign Keys) when suggesting analysis types.
    
    Provide your response in a clear, structured format.
    """
    
    # Generate response using LLM
    with st.spinner("Analyzing the database..."):
        response = await llm_wrapper.generate_response_async(prompt)
        st.session_state[f"{page_key}_llm_response"] = response  # Save the response in session state

        # Save the report in DOCX format
        report_buffer = save_report_to_docx(response)
        st.session_state[f"{page_key}_report"] = report_buffer  # Save the report in session state

# Connect to the database
connection = connect_to_database()
if connection:
    st.success("Database connection established successfully.")

    # Analyze the database and provide recommendations
    if st.button("Analyze Database"):
        asyncio.run(analyze_database(connection))
        
    # Display the LLM response, if it exists in the session state
    if st.session_state[f"{page_key}_llm_response"]:
        st.write("## Analysis Report")
        st.write(st.session_state[f"{page_key}_llm_response"])
        
        # Provide download link for the report in DOCX format
        if st.session_state[f"{page_key}_report"]:
            st.download_button(
                label="Download Report",
                data=st.session_state[f"{page_key}_report"],
                file_name="database_analysis_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )